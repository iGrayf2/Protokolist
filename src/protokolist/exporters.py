from __future__ import annotations

from pathlib import Path

from docx import Document

from .segments import TranscriptSegment, format_srt_timestamp, format_timestamp


def write_txt(segments: list[TranscriptSegment], path: str | Path) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as file:
        for seg in segments:
            file.write(f"[{format_timestamp(seg.start)}] {seg.text}\n")
    return path


def write_srt(segments: list[TranscriptSegment], path: str | Path) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as file:
        for index, seg in enumerate(segments, start=1):
            file.write(f"{index}\n")
            file.write(f"{format_srt_timestamp(seg.start)} --> {format_srt_timestamp(seg.end)}\n")
            file.write(f"{seg.text}\n\n")
    return path


def write_docx(segments: list[TranscriptSegment], path: str | Path, title: str = "Стенограмма совещания") -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_heading(title, level=1)
    for seg in segments:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[{format_timestamp(seg.start)}] ").bold = True
        paragraph.add_run(seg.text)
    document.save(path)
    return path
